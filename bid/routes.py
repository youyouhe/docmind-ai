"""
Bid writing routes for PageIndex API.

Provides endpoints for:
- Project CRUD operations
- Auto-save functionality
- AI content generation
- Text rewriting
"""

import asyncio
import io
import logging
import os
import sys
import uuid
import json
from typing import List, Optional
from datetime import datetime

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from api.database import get_db, DatabaseManager
from api.services import LLMProvider
from api.websocket_manager import manager

logger = logging.getLogger(__name__)


router = APIRouter(
    prefix="/api/bid",
    tags=["bid-writing"],
)


# =============================================================================
# Data Models
# =============================================================================

class TenderSection(BaseModel):
    """A section in the bid document."""
    id: str
    title: str
    content: str
    summary: Optional[str] = None
    requirement_references: List[str] = []
    status: str = "pending"  # pending/in_progress/completed
    order: int
    word_count: int = 0


class TenderProject(BaseModel):
    """A bid writing project."""
    id: str
    title: str
    tender_document_id: str
    tender_document_tree: dict
    sections: List[TenderSection]
    status: str = "draft"  # draft/review/completed
    version: int = 1
    created_at: int
    updated_at: int


class CreateProjectRequest(BaseModel):
    """Request to create a new project."""
    title: str
    tender_document_id: str
    tender_document_tree: dict
    sections: List[TenderSection]


class GenerateContentRequest(BaseModel):
    """Request to generate section content."""
    section_id: str
    section_title: str
    section_description: str
    tender_tree: dict
    requirement_references: List[str] = []
    previous_context: Optional[str] = None
    user_prompt: Optional[str] = None
    attachments: Optional[List[str]] = None


class RewriteRequest(BaseModel):
    """Request to rewrite text."""
    text: str
    mode: str  # formal/concise/expand/clarify
    context: Optional[str] = None


class AutoSaveRequest(BaseModel):
    """Request to auto-save section content."""
    content: str


class ExportRequest(BaseModel):
    """Request to export a project as Word/PDF."""
    format: str = "word"  # word or pdf
    include_outline: bool = True
    include_requirements: bool = False


class GenerateOutlineRequest(BaseModel):
    """Request to generate outline via multi-agent pipeline."""
    tender_document_id: str
    tender_document_tree: dict
    title: str = "新投标项目"
    user_requirements: Optional[str] = None
    attachment_names: Optional[List[str]] = None


class WriteContentRequest(BaseModel):
    """Request to write section content via agents."""
    section_ids: Optional[List[str]] = None  # None = all pending sections


class ReviewRequest(BaseModel):
    """Request to review bid document via agents."""
    pass


# =============================================================================
# Agent-based Outline Generation
# =============================================================================

@router.post("/outline/generate")
async def generate_outline_via_agents(request: GenerateOutlineRequest) -> dict:
    """Start the multi-agent outline generation pipeline.

    Creates a draft project, launches format-extractor → outline-planner
    in the background, and returns immediately with the project ID.
    The frontend should subscribe to WebSocket ``audit_progress`` and
    ``status_update`` messages using the returned project ID.
    """
    db = get_db()
    project_id = f"project-{uuid.uuid4()}"

    # Create draft project (empty sections — agents will fill them)
    db.create_bid_project(
        project_id=project_id,
        title=request.title,
        tender_document_id=request.tender_document_id,
        tender_document_tree=request.tender_document_tree,
        sections=[],
    )

    # Launch pipeline in background
    asyncio.create_task(
        _run_outline_pipeline_task(
            project_id=project_id,
            tender_document_id=request.tender_document_id,
            tender_document_tree=request.tender_document_tree,
            title=request.title,
            user_requirements=request.user_requirements,
            attachment_names=request.attachment_names,
        )
    )

    return {"project_id": project_id, "status": "started"}


async def _run_outline_pipeline_task(
    project_id: str,
    tender_document_id: str,
    tender_document_tree: dict,
    title: str,
    user_requirements: Optional[str] = None,
    attachment_names: Optional[List[str]] = None,
) -> None:
    """Background task: run the agent pipeline and broadcast results via WS."""
    # Ensure bid_agents is importable (repo root may not be on sys.path)
    repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    if repo_root not in sys.path:
        sys.path.insert(0, repo_root)

    from bid_agents.outline_pipeline import run_outline_pipeline
    logger.info("Outline pipeline task started for project %s", project_id)

    api_url = f"http://localhost:{os.getenv('PORT', '8003')}"

    async def progress_callback(phase: str, message: str) -> None:
        phase_map = {"format_extraction": (1, 2), "outline_planning": (2, 2)}
        phase_number, total_phases = phase_map.get(phase, (1, 2))
        progress = (phase_number - 1) / total_phases * 100
        await manager.broadcast_audit_progress(
            document_id=project_id,
            phase=phase,
            phase_number=phase_number,
            total_phases=total_phases,
            message=message,
            progress=progress,
        )

    try:
        logger.info("[outline-task] Calling run_outline_pipeline for %s", project_id)
        sections = await asyncio.wait_for(
            run_outline_pipeline(
                project_id=project_id,
                api_url=api_url,
                user_requirements=user_requirements,
                attachment_names=attachment_names,
                progress_callback=progress_callback,
            ),
            timeout=600,  # 10 min — each RAG query takes ~30s
        )
        logger.info("[outline-task] Pipeline complete! %d sections for %s", len(sections), project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="completed",
            metadata={"outline": sections},
        )
    except asyncio.TimeoutError:
        logger.error("[outline-task] Timed out for project %s", project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message="大纲生成超时（600秒）",
        )
    except Exception as e:
        logger.exception("[outline-task] FAILED for project %s: %s", project_id, e)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message=str(e),
        )


# =============================================================================
# Agent-based Content Writing
# =============================================================================

@router.post("/projects/{project_id}/content/write")
async def write_content_via_agents(project_id: str, request: WriteContentRequest) -> dict:
    """Start the multi-agent content writing pipeline.

    Launches writer agents (commercial / technical / pricing) in the background
    to write section content.  Broadcasts progress via WebSocket.
    """
    db = get_db()
    project = db.get_bid_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")

    asyncio.create_task(
        _run_content_pipeline_task(
            project_id=project_id,
            section_ids=request.section_ids,
        )
    )

    return {"project_id": project_id, "status": "started"}


async def _run_content_pipeline_task(
    project_id: str,
    section_ids: Optional[List[str]] = None,
) -> None:
    """Background task: run the content writing pipeline."""
    repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    if repo_root not in sys.path:
        sys.path.insert(0, repo_root)

    from bid_agents.content_pipeline import run_content_pipeline

    api_url = f"http://localhost:{os.getenv('PORT', '8003')}"

    async def progress_callback(phase: str, message: str, current: int, total: int) -> None:
        progress = current / total * 100 if total > 0 else 0
        await manager.broadcast_audit_progress(
            document_id=project_id,
            phase=phase,
            phase_number=current,
            total_phases=total,
            message=message,
            progress=progress,
        )

    try:
        result = await asyncio.wait_for(
            run_content_pipeline(
                project_id=project_id,
                api_url=api_url,
                section_ids=section_ids,
                progress_callback=progress_callback,
            ),
            timeout=600,  # 10 min for writing multiple sections
        )
        await manager.broadcast_status_update(
            document_id=project_id,
            status="completed",
            metadata={
                "pipeline": "content",
                "written": result["written"],
                "failed": result["failed"],
                "sections": result["sections"],
            },
        )
    except asyncio.TimeoutError:
        logger.error("Content pipeline timed out for project %s", project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message="内容编写超时（600秒）",
        )
    except Exception as e:
        logger.exception("Content pipeline failed for project %s", project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message=str(e),
        )


# =============================================================================
# Agent-based Review
# =============================================================================

@router.post("/projects/{project_id}/review")
async def review_via_agents(project_id: str, request: ReviewRequest) -> dict:
    """Start the review + compliance check pipeline.

    Launches review-agent → compliance-checker in the background.
    Broadcasts progress via WebSocket.
    """
    db = get_db()
    project = db.get_bid_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")

    asyncio.create_task(
        _run_review_pipeline_task(project_id=project_id)
    )

    return {"project_id": project_id, "status": "started"}


async def _run_review_pipeline_task(project_id: str) -> None:
    """Background task: run the review pipeline."""
    repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    if repo_root not in sys.path:
        sys.path.insert(0, repo_root)

    from bid_agents.review_pipeline import run_review_pipeline

    api_url = f"http://localhost:{os.getenv('PORT', '8003')}"

    async def progress_callback(phase: str, message: str) -> None:
        phase_map = {"quality_review": (1, 2), "compliance_check": (2, 2)}
        phase_number, total_phases = phase_map.get(phase, (1, 2))
        progress = (phase_number - 1) / total_phases * 100
        await manager.broadcast_audit_progress(
            document_id=project_id,
            phase=phase,
            phase_number=phase_number,
            total_phases=total_phases,
            message=message,
            progress=progress,
        )

    try:
        result = await asyncio.wait_for(
            run_review_pipeline(
                project_id=project_id,
                api_url=api_url,
                progress_callback=progress_callback,
            ),
            timeout=300,  # 5 min for review
        )
        await manager.broadcast_status_update(
            document_id=project_id,
            status="completed",
            metadata={
                "pipeline": "review",
                "review_feedback": result["review_feedback"],
                "compliance_matrix": result["compliance_matrix"],
                "summary": result["summary"],
            },
        )
    except asyncio.TimeoutError:
        logger.error("Review pipeline timed out for project %s", project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message="审核超时（300秒）",
        )
    except Exception as e:
        logger.exception("Review pipeline failed for project %s", project_id)
        await manager.broadcast_status_update(
            document_id=project_id,
            status="failed",
            error_message=str(e),
        )


# =============================================================================
# Project Endpoints (Database-backed)
# =============================================================================

@router.post("/projects")
async def create_project(request: CreateProjectRequest) -> dict:
    """Create a new bid writing project."""
    db = get_db()
    project_id = f"project-{uuid.uuid4()}"

    sections = [s.model_dump() for s in request.sections]
    result = db.create_bid_project(
        project_id=project_id,
        title=request.title,
        tender_document_id=request.tender_document_id,
        tender_document_tree=request.tender_document_tree,
        sections=sections,
    )
    return result


@router.get("/projects")
async def list_projects() -> list:
    """List all bid writing projects."""
    db = get_db()
    return db.list_bid_projects()


@router.get("/projects/{project_id}")
async def get_project(project_id: str) -> dict:
    """Get a specific bid writing project."""
    db = get_db()
    project = db.get_bid_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")
    return project


@router.put("/projects/{project_id}")
async def update_project(project_id: str, request: TenderProject) -> dict:
    """Update a bid writing project."""
    db = get_db()
    sections = [s.model_dump() for s in request.sections]
    result = db.update_bid_project(
        project_id=project_id,
        title=request.title,
        status=request.status,
        sections=sections,
        tender_document_tree=request.tender_document_tree,
    )
    if not result:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")
    return result


@router.delete("/projects/{project_id}")
async def delete_project(project_id: str) -> dict:
    """Delete a bid writing project."""
    db = get_db()
    deleted = db.delete_bid_project(project_id)
    if not deleted:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")
    return {"id": project_id, "deleted": True}


@router.post("/projects/{project_id}/sections/{section_id}/auto-save")
async def auto_save_section(
    project_id: str,
    section_id: str,
    request: AutoSaveRequest
) -> dict:
    """Auto-save a section's content."""
    db = get_db()
    updated_at = db.auto_save_bid_section(project_id, section_id, request.content)
    if updated_at is None:
        raise HTTPException(status_code=404, detail=f"Project or section not found")
    return {"success": True, "saved_at": updated_at}


# =============================================================================
# Export Endpoint
# =============================================================================

@router.post("/projects/{project_id}/export")
async def export_project(project_id: str, request: ExportRequest):
    """Export a bid writing project as Word document."""
    db = get_db()
    project_data = db.get_bid_project(project_id)
    if not project_data:
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")

    project = TenderProject(**project_data)

    if request.format == "word":
        return _export_to_word(project, request)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported export format: {request.format}. Currently only 'word' is supported.")


def _export_to_word(project: TenderProject, request: ExportRequest) -> StreamingResponse:
    """Generate a Word document from the project."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title page
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project.title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(datetime.fromtimestamp(project.created_at / 1000).strftime('%Y年%m月%d日'))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Table of Contents (outline)
    if request.include_outline:
        toc_heading = doc.add_heading('目录', level=1)
        toc_heading.runs[0].font.size = Pt(16)

        for i, section in enumerate(sorted(project.sections, key=lambda s: s.order)):
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.space_after = Pt(4)
            run = toc_para.add_run(f"{i + 1}. {section.title}")
            run.font.size = Pt(11)
            if section.status == "completed":
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.color.rgb = RGBColor(180, 180, 180)

        doc.add_page_break()

    # Section content
    for i, section in enumerate(sorted(project.sections, key=lambda s: s.order)):
        # Section heading
        heading = doc.add_heading(f"{i + 1}. {section.title}", level=1)
        heading.runs[0].font.size = Pt(16)

        # Requirement references
        if request.include_requirements and section.summary:
            req_para = doc.add_paragraph()
            req_para.paragraph_format.space_after = Pt(8)
            run = req_para.add_run(f"【招标要求摘要】{section.summary}")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Section body
        if section.content.strip():
            for paragraph_text in section.content.split('\n'):
                if not paragraph_text.strip():
                    doc.add_paragraph()
                    continue

                # Handle markdown-like headers in content
                if paragraph_text.startswith('### '):
                    h = doc.add_heading(paragraph_text[4:], level=3)
                    h.runs[0].font.size = Pt(12)
                elif paragraph_text.startswith('## '):
                    h = doc.add_heading(paragraph_text[3:], level=2)
                    h.runs[0].font.size = Pt(14)
                else:
                    p = doc.add_paragraph(paragraph_text)
                    p.paragraph_format.line_spacing = Pt(22)
                    for run in p.runs:
                        run.font.size = Pt(11)
        else:
            p = doc.add_paragraph('（此章节尚未编写）')
            p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
            p.runs[0].font.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{project.title}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


# =============================================================================
# Content Generation Endpoints
# =============================================================================

@router.post("/content/generate")
async def generate_content(request: GenerateContentRequest) -> dict:
    """Generate bid section content using AI."""
    from api.services import ParseService

    # Get LLM provider
    provider = os.getenv("LLM_PROVIDER", "deepseek")
    api_key = os.getenv(f"{provider.upper()}_API_KEY")
    model = os.getenv("LLM_MODEL", None)

    llm = LLMProvider(provider=provider, api_key=api_key, model=model)

    # Build prompt
    prompt = build_content_generation_prompt(
        request.section_title,
        request.section_description,
        request.tender_tree,
        request.requirement_references,
        request.previous_context,
        request.user_prompt,
        request.attachments
    )

    # Generate content
    content = await llm.chat(prompt)

    return {
        "content": content,
        "provider": provider,
        "model": model or llm.get_default_model(),
        "generated_at": int(datetime.now().timestamp() * 1000)
    }


@router.post("/content/rewrite")
async def rewrite_text(request: RewriteRequest) -> dict:
    """Rewrite text using AI."""
    from api.services import ParseService

    # Get LLM provider
    provider = os.getenv("LLM_PROVIDER", "deepseek")
    api_key = os.getenv(f"{provider.upper()}_API_KEY")
    model = os.getenv("LLM_MODEL", None)

    llm = LLMProvider(provider=provider, api_key=api_key, model=model)

    # Build prompt
    prompt = build_rewrite_prompt(request.text, request.mode, request.context)

    # Rewrite text
    rewritten_text = await llm.chat(prompt)

    return {
        "rewritten_text": rewritten_text,
        "provider": provider,
        "model": model or llm.get_default_model()
    }


# =============================================================================
# Helper Functions
# =============================================================================

def build_content_generation_prompt(
    section_title: str,
    section_description: str,
    tender_tree: dict,
    requirement_references: List[str],
    previous_context: Optional[str],
    user_prompt: Optional[str],
    attachments: Optional[List[str]]
) -> str:
    """Build prompt for AI content generation."""

    # Extract relevant content from tender tree
    def extract_tree_content(node: dict, depth: int = 0, max_depth: int = 2) -> List[str]:
        if depth > max_depth:
            return []

        content_parts = []

        title = node.get("title", "")
        summary = node.get("summary", "")

        node_text = f"【{title}】"
        if summary:
            node_text += f"\n{summary}"

        content_parts.append(node_text)

        for child in node.get("children", []):
            content_parts.extend(extract_tree_content(child, depth + 1, max_depth))

        return content_parts

    # Build tender requirements context
    tender_context = ""
    if requirement_references:
        context_parts = extract_tree_content(tender_tree)
        tender_context = "\n\n".join(context_parts[:10])

    prompt = f"""你是一个专业的投标文件编写专家。请根据招标要求编写投标文件章节。

=== 当前章节 ===
标题：{section_title}
描述：{section_description}

"""

    if tender_context:
        prompt += f"=== 招标文档相关章节 ===\n{tender_context}\n\n"

    if requirement_references:
        prompt += f"=== 对应招标要求节点 ===\n"
        prompt += f"引用: {', '.join(requirement_references)}\n\n"

    if previous_context:
        prompt += f"=== 前文上下文 ===\n{previous_context[:1000]}...\n\n"

    if user_prompt:
        prompt += f"=== 用户补充说明 ===\n{user_prompt}\n\n"

    if attachments:
        prompt += f"=== 参考附件 ===\n{', '.join(attachments)}\n\n"

    prompt += """=== 编写要求 ===
1. 严格按照招标要求编写
2. 使用专业、规范的商务语言
3. 突出公司优势和符合性
4. 结构清晰，逻辑严密
5. 语言简练，避免冗余

请生成该章节的完整内容。直接返回内容，不要包含任何解释或元信息。"""

    return prompt


def build_rewrite_prompt(
    original_text: str,
    mode: str,
    context: Optional[str]
) -> str:
    """Build prompt for text rewriting."""

    mode_descriptions = {
        'formal': '正式化 - 使语言更加正式、规范，符合商务文件要求',
        'concise': '精简 - 去除冗余，保留核心信息，使表达更加简洁',
        'expand': '扩充 - 在原有内容基础上增加详细说明和补充信息',
        'clarify': '澄清 - 使表达更加清晰明确，避免歧义'
    }

    mode_instructions = {
        'formal': '使用商务正式用语，适当使用被动语态，增加专业术语',
        'concise': '删除重复和不必要的词语，保留关键信息，使用更直接的表达',
        'expand': '在原有内容基础上增加详细说明、具体例子和补充信息',
        'clarify': '重新组织句子结构，使用更明确的词汇，消除模糊表达'
    }

    description = mode_descriptions.get(mode, 'formal')
    instruction = mode_instructions.get(mode, '使用商务正式用语')

    prompt = f"""请对以下投标文件文本进行改写。

=== 改写模式 ===
{description}

=== 原文 ===
{original_text}

"""

    if context:
        prompt += f"=== 上下文 ===\n{context}\n\n"

    prompt += f"""=== 改写要求 ===
{instruction}

请直接返回改写后的文本，不要解释或添加其他内容。"""

    return prompt
